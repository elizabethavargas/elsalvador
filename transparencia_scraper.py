"""
transparencia_scraper.py — Scrape document metadata AND full text
from https://www.transparencia.gob.sv/ for all publications since 2015.

SITE ARCHITECTURE:
  indexJSON.php
    → list of 11 institution category tokens

  institucionesCategoriaJSON.php?id_tipo={N}
    → JSON array of institutions in that category
    N=1  Ministerios (17)       N=6  Hospitales (36)
    N=2  Autónomas (97)         N=7  Gobernaciones (14)
    N=3  Presidencia (1)        N=8  ONG (2)
    N=4  Otras dependencias     N=10 Nuevas Municipalidades (41)
    N=5  Municipalidades (231)  N=11 Órgano Legislativo (1+)

  perfilInstitucionesDocMasDescargadosJson.php?id_institucion={id}
    → JSON {data:[...]} — most-downloaded documents per institution.
      Fields: id_documents, downloads, active, year, estandar, name,
              document_file_name_anexo, file_url, file_url2, file_detalle

  descarga_archivo.php?id={base64(doc_id)}&inst={doc_id}
    → actual file download (PDF, XLSX, DOCX, …)

DEPENDENCIES (pip install):
    requests beautifulsoup4 lxml
    pdfminer.six          ← PDF text extraction
    python-docx           ← DOCX text extraction
    openpyxl              ← XLSX text extraction
    xlrd                  ← XLS  text extraction

USAGE:
    # Metadata only (fast — no file downloads)
    python transparencia_scraper.py

    # Download files AND extract their text into the CSV/JSONL
    python transparencia_scraper.py --extract-text

    # Save the raw files to disk as well
    python transparencia_scraper.py --extract-text --download

    # Only Ministerios
    python transparencia_scraper.py --extract-text --id-tipo 1

    # Single institution (e.g. Presidencia id=9)
    python transparencia_scraper.py --extract-text --institution 9

    # Resume an interrupted run
    python transparencia_scraper.py --extract-text --resume

OUTPUT (in ./output/transparencia/):
    metadata.csv          — one row per document (includes `text` when --extract-text)
    metadata.jsonl        — same, one JSON object per line
    checkpoint.json       — progress state for --resume
    files/{inst}/{year}/  — raw downloaded files (only with --download)
"""

import argparse
import base64
import io
import json
import logging
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple
from urllib.parse import urljoin, urlparse

import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# ─────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────
BASE_URL = "https://www.transparencia.gob.sv/"
START_YEAR = 2015
MAX_ID_TIPO = 20

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "es-SV,es;q=0.9,en;q=0.5",
    "Accept": "application/json, text/html, */*",
    "Referer": BASE_URL,
}

REQUEST_TIMEOUT = 60   # longer for file downloads
MAX_RETRIES = 3

_MIME_TO_EXT = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xls",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "application/zip": "zip",
    "text/html": "html",
    "text/plain": "txt",
    "image/jpeg": "jpg",
    "image/png": "png",
}

# Fields written to CSV/JSONL — text column only present when --extract-text
CSV_FIELDS_BASE = [
    "id_documents", "id_institucion", "institution_name", "institution_acronym",
    "id_tipo", "year", "estandar", "name", "active", "downloads",
    "file_url", "file_url2", "file_detalle", "has_annex",
    "file_type", "file_extension", "text_length",
    "local_path", "scraped_at",
]
CSV_FIELDS_WITH_TEXT = CSV_FIELDS_BASE + ["text"]


# ─────────────────────────────────────────────────────────────
# LOGGING
# ─────────────────────────────────────────────────────────────
logger = logging.getLogger("transparencia")


def _setup_logging(log_dir: Path) -> None:
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / f"transparencia_{datetime.now():%Y%m%d_%H%M%S}.log"
    fmt = "%(asctime)s [%(levelname)s] %(message)s"
    logging.basicConfig(
        level=logging.INFO,
        format=fmt,
        handlers=[
            logging.StreamHandler(sys.stdout),
            logging.FileHandler(log_file, encoding="utf-8"),
        ],
    )


# ─────────────────────────────────────────────────────────────
# HTTP SESSION
# ─────────────────────────────────────────────────────────────
_session: Optional[requests.Session] = None
_last_request_time: float = 0.0
_request_delay: float = 1.5


def _build_session() -> requests.Session:
    session = requests.Session()
    session.headers.update(HEADERS)
    retry = Retry(
        total=MAX_RETRIES,
        backoff_factor=2,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET", "HEAD"],
    )
    adapter = HTTPAdapter(max_retries=retry)
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    return session


def _get(url: str, delay: Optional[float] = None, **kwargs) -> Optional[requests.Response]:
    """Rate-limited GET. Returns None on any failure."""
    global _last_request_time
    used_delay = delay if delay is not None else _request_delay
    elapsed = time.time() - _last_request_time
    if elapsed < used_delay:
        time.sleep(used_delay - elapsed)
    try:
        resp = _session.get(url, timeout=REQUEST_TIMEOUT, **kwargs)  # type: ignore[union-attr]
        _last_request_time = time.time()
        if resp.status_code == 200:
            return resp
        logger.warning("HTTP %d: %s", resp.status_code, url)
        return None
    except requests.RequestException as exc:
        logger.warning("Request error %s: %s", url, exc)
        return None


# ─────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────
def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF bytes using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(data))
        return _clean_text(text)
    except ImportError:
        logger.warning("pdfminer.six not installed — run: pip install pdfminer.six")
        return ""
    except Exception as exc:
        logger.debug("PDF extraction error: %s", exc)
        return ""


def _extract_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return _clean_text("\n".join(parts))
    except ImportError:
        logger.warning("python-docx not installed — run: pip install python-docx")
        return ""
    except Exception as exc:
        logger.debug("DOCX extraction error: %s", exc)
        return ""


def _extract_xlsx(data: bytes) -> str:
    """Extract text from XLSX bytes using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
        return _clean_text("\n".join(parts))
    except ImportError:
        logger.warning("openpyxl not installed — run: pip install openpyxl")
        return ""
    except Exception as exc:
        logger.debug("XLSX extraction error: %s", exc)
        return ""


def _extract_xls(data: bytes) -> str:
    """Extract text from XLS bytes using xlrd."""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        parts = []
        for sheet in wb.sheets():
            for rx in range(sheet.nrows):
                cells = [str(sheet.cell(rx, cx).value).strip()
                         for cx in range(sheet.ncols)
                         if str(sheet.cell(rx, cx).value).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _clean_text("\n".join(parts))
    except ImportError:
        logger.warning("xlrd not installed — run: pip install xlrd")
        return ""
    except Exception as exc:
        logger.debug("XLS extraction error: %s", exc)
        return ""


def _extract_html(data: bytes) -> str:
    """Extract text from HTML bytes using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(data, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return _clean_text(soup.get_text(separator="\n"))
    except Exception as exc:
        logger.debug("HTML extraction error: %s", exc)
        return ""


def _clean_text(text: str) -> str:
    """Normalise whitespace and strip boilerplate."""
    if not text:
        return ""
    # Collapse runs of blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse horizontal whitespace
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def extract_text_from_bytes(data: bytes, content_type: str) -> Tuple[str, str]:
    """
    Dispatch text extraction based on MIME type.
    Returns (file_type, extracted_text).
    """
    ct = content_type.lower().split(";")[0].strip()

    if ct == "application/pdf":
        return "pdf", _extract_pdf(data)

    if ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
              "application/msword"):
        return "docx", _extract_docx(data)

    if ct == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return "xlsx", _extract_xlsx(data)

    if ct == "application/vnd.ms-excel":
        return "xls", _extract_xls(data)

    if ct in ("text/html", "application/xhtml+xml"):
        return "html", _extract_html(data)

    if ct == "text/plain":
        try:
            return "txt", _clean_text(data.decode("utf-8", errors="replace"))
        except Exception:
            return "txt", ""

    # Unknown type — try PDF first (many government portals serve PDF without correct MIME)
    if data[:4] == b"%PDF":
        return "pdf", _extract_pdf(data)

    # Try ZIP (some docs come as ZIP of XMLs e.g. XLSX/DOCX)
    if data[:2] == b"PK":
        # Could be DOCX or XLSX — attempt both
        text = _extract_docx(data)
        if text:
            return "docx", text
        text = _extract_xlsx(data)
        if text:
            return "xlsx", text

    logger.debug("Unsupported content-type for text extraction: %s", ct)
    return ct, ""


# ─────────────────────────────────────────────────────────────
# CHECKPOINT
# ─────────────────────────────────────────────────────────────
class Checkpoint:
    def __init__(self, path: Path):
        self.path = path
        self.data: dict = {
            "completed_institutions": [],
            "total_docs": 0,
            "total_extracted": 0,
            "started_at": datetime.now().isoformat(),
        }
        if path.exists():
            try:
                with open(path, encoding="utf-8") as f:
                    self.data = json.load(f)
                logger.info("Resumed: %d institutions already done",
                            len(self.data["completed_institutions"]))
            except (json.JSONDecodeError, KeyError):
                logger.warning("Corrupt checkpoint — starting fresh")

    def save(self) -> None:
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(self.data, f, ensure_ascii=False, indent=2)

    def mark_done(self, inst_id: int, doc_count: int, extracted: int) -> None:
        if inst_id not in self.data["completed_institutions"]:
            self.data["completed_institutions"].append(inst_id)
        self.data["total_docs"] += doc_count
        self.data["total_extracted"] += extracted
        self.save()

    def is_done(self, inst_id: int) -> bool:
        return inst_id in self.data["completed_institutions"]


# ─────────────────────────────────────────────────────────────
# INSTITUTION DISCOVERY
# ─────────────────────────────────────────────────────────────
def fetch_institutions_for_tipo(id_tipo: int) -> list:
    url = urljoin(BASE_URL, f"institucionesCategoriaJSON.php?id_tipo={id_tipo}")
    resp = _get(url)
    if resp is None:
        return []
    try:
        payload = resp.json()
        insts = payload.get("data", [])
        return insts if isinstance(insts, list) else []
    except ValueError:
        return []


def discover_all_institutions(only_id_tipo: Optional[int] = None) -> list:
    all_institutions: list = []
    seen_ids: set = set()
    probe_range = [only_id_tipo] if only_id_tipo else range(1, MAX_ID_TIPO + 1)

    for id_tipo in probe_range:
        insts = fetch_institutions_for_tipo(id_tipo)
        if not insts:
            continue
        logger.info("id_tipo=%d: %d institutions", id_tipo, len(insts))
        for inst in insts:
            try:
                inst_id = int(inst.get("id_institucion", 0))
            except (TypeError, ValueError):
                continue
            if inst_id == 0 or inst_id in seen_ids:
                continue
            seen_ids.add(inst_id)
            all_institutions.append({
                "id_institucion": inst_id,
                "nombre_institucion": inst.get("nombre_institucion", ""),
                "acronym": inst.get("acronym", ""),
                "id_tipo": id_tipo,
            })

    logger.info("Total unique institutions: %d", len(all_institutions))
    return all_institutions


# ─────────────────────────────────────────────────────────────
# DOCUMENT DISCOVERY
# ─────────────────────────────────────────────────────────────
def fetch_documents_for_institution(inst_id: int) -> list:
    url = urljoin(
        BASE_URL,
        f"perfilInstitucionesDocMasDescargadosJson.php?id_institucion={inst_id}",
    )
    resp = _get(url)
    if resp is None:
        return []
    try:
        payload = resp.json()
        if isinstance(payload, list):
            return payload
        return payload.get("data", []) or []
    except ValueError:
        return []


def _safe_int(val, default: int = 0) -> int:
    try:
        if val in (None, "", "null"):
            return default
        return int(str(val).replace(",", ""))
    except (TypeError, ValueError):
        return default


def _safe_year(val) -> Optional[int]:
    try:
        y = int(val)
        return y if 1990 <= y <= 2100 else None
    except (TypeError, ValueError):
        return None


def _make_download_url(doc_id: int) -> str:
    b64 = base64.b64encode(str(doc_id).encode()).decode()
    return urljoin(BASE_URL, f"descarga_archivo.php?id={b64}&inst={doc_id}")


def _infer_extension(url: str) -> str:
    try:
        _, ext = os.path.splitext(urlparse(url).path)
        return ext.lower().lstrip(".")
    except Exception:
        return ""


def build_record(raw: dict, institution: dict) -> Optional[dict]:
    doc_id = _safe_int(raw.get("id_documents"))
    if doc_id == 0:
        return None
    year = _safe_year(raw.get("year"))
    if year is not None and year < START_YEAR:
        return None

    file_url = raw.get("file_url") or _make_download_url(doc_id)
    return {
        "id_documents": doc_id,
        "id_institucion": institution["id_institucion"],
        "institution_name": institution["nombre_institucion"],
        "institution_acronym": institution["acronym"],
        "id_tipo": institution["id_tipo"],
        "year": year,
        "estandar": raw.get("estandar", ""),
        "name": (raw.get("name") or "")[:500],
        "active": raw.get("active", ""),
        "downloads": _safe_int(raw.get("downloads")),
        "file_url": file_url,
        "file_url2": raw.get("file_url2") or "",
        "file_detalle": raw.get("file_detalle") or "",
        "has_annex": bool(raw.get("document_file_name_anexo")),
        "file_type": "",
        "file_extension": _infer_extension(file_url),
        "text": "",
        "text_length": 0,
        "local_path": "",
        "scraped_at": datetime.now().isoformat(),
    }


# ─────────────────────────────────────────────────────────────
# DOWNLOAD + EXTRACT
# ─────────────────────────────────────────────────────────────
def _safe_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]', "_", name)
    return re.sub(r"\s+", "_", name.strip())[:200] or "document"


def fetch_and_extract(record: dict, files_dir: Optional[Path]) -> bool:
    """
    Download the file for a record, extract its text, and optionally save to disk.
    Mutates the record in place with: text, file_type, file_extension, text_length, local_path.
    Returns True if text was successfully extracted.
    """
    url = record["file_url"]
    if not url:
        return False

    resp = _get(url, delay=_request_delay * 1.3)
    if resp is None:
        return False

    data = resp.content
    ctype = resp.headers.get("Content-Type", "application/octet-stream").split(";")[0].strip()

    # Resolve extension from actual Content-Type
    ext = _MIME_TO_EXT.get(ctype.lower(), record["file_extension"] or "bin")
    record["file_extension"] = ext

    # Extract text
    file_type, text = extract_text_from_bytes(data, ctype)
    record["file_type"] = file_type
    record["text"] = text
    record["text_length"] = len(text)

    # Optionally save raw file
    if files_dir is not None:
        inst_slug = _safe_filename(record["institution_acronym"] or str(record["id_institucion"]))
        year_str = str(record["year"]) if record["year"] else "unknown_year"
        dest_dir = files_dir / inst_slug / year_str
        dest_dir.mkdir(parents=True, exist_ok=True)
        doc_name = _safe_filename(record["name"] or str(record["id_documents"]))
        dest_path = dest_dir / f"{record['id_documents']}_{doc_name}.{ext}"
        if not dest_path.exists():
            try:
                dest_path.write_bytes(data)
            except OSError as exc:
                logger.warning("Write error %s: %s", dest_path, exc)
        record["local_path"] = str(dest_path)

    return bool(text)


# ─────────────────────────────────────────────────────────────
# OUTPUT WRITERS
# ─────────────────────────────────────────────────────────────
def _init_csv(csv_path: Path, with_text: bool) -> None:
    if not csv_path.exists():
        fields = CSV_FIELDS_WITH_TEXT if with_text else CSV_FIELDS_BASE
        with open(csv_path, "w", encoding="utf-8") as f:
            f.write(",".join(fields) + "\n")


def _append_csv(csv_path: Path, records: list, with_text: bool) -> None:
    fields = CSV_FIELDS_WITH_TEXT if with_text else CSV_FIELDS_BASE
    with open(csv_path, "a", encoding="utf-8", newline="") as f:
        for rec in records:
            row = []
            for field in fields:
                val = str(rec.get(field, "")).replace('"', '""')
                if any(ch in val for ch in (",", '"', "\n", "\r")):
                    val = f'"{val}"'
                row.append(val)
            f.write(",".join(row) + "\n")


def _append_jsonl(jsonl_path: Path, records: list, with_text: bool) -> None:
    fields = CSV_FIELDS_WITH_TEXT if with_text else CSV_FIELDS_BASE
    with open(jsonl_path, "a", encoding="utf-8") as f:
        for rec in records:
            out = {k: rec.get(k, "") for k in fields}
            f.write(json.dumps(out, ensure_ascii=False) + "\n")


# ─────────────────────────────────────────────────────────────
# MAIN SCRAPER
# ─────────────────────────────────────────────────────────────
def run(
    only_id_tipo: Optional[int] = None,
    only_institution: Optional[int] = None,
    extract_text: bool = False,
    do_download: bool = False,
    resume: bool = False,
    output_dir: Path = Path("output/transparencia"),
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    files_dir: Optional[Path] = None
    if do_download:
        files_dir = output_dir / "files"
        files_dir.mkdir(parents=True, exist_ok=True)

    csv_path = output_dir / "metadata.csv"
    jsonl_path = output_dir / "metadata.jsonl"
    ckpt_path = output_dir / "checkpoint.json"

    _init_csv(csv_path, with_text=extract_text)
    checkpoint = Checkpoint(ckpt_path)

    if not resume:
        checkpoint.data["completed_institutions"] = []
        checkpoint.save()

    # ── Discover institutions ──────────────────────────────────
    if only_institution:
        institutions = [{
            "id_institucion": only_institution,
            "nombre_institucion": f"Institution {only_institution}",
            "acronym": str(only_institution),
            "id_tipo": 0,
        }]
    else:
        institutions = discover_all_institutions(only_id_tipo=only_id_tipo)

    if not institutions:
        logger.error("No institutions found — exiting.")
        return

    logger.info(
        "Scraping %d institutions | year >= %d | extract_text=%s | download=%s",
        len(institutions), START_YEAR, extract_text, do_download,
    )

    total_docs = 0
    total_extracted = 0

    # ── Documents per institution ──────────────────────────────
    for i, inst in enumerate(institutions, 1):
        inst_id = inst["id_institucion"]

        if checkpoint.is_done(inst_id):
            logger.info("[%d/%d] Skip (done): %s [id=%d]",
                        i, len(institutions), inst["nombre_institucion"], inst_id)
            continue

        logger.info("[%d/%d] %s [id=%d]",
                    i, len(institutions), inst["nombre_institucion"], inst_id)

        raw_docs = fetch_documents_for_institution(inst_id)
        records: list = []
        for raw in raw_docs:
            rec = build_record(raw, inst)
            if rec is not None:
                records.append(rec)

        logger.info("  %d raw → %d kept (year >= %d)",
                    len(raw_docs), len(records), START_YEAR)

        extracted_count = 0
        if extract_text and records:
            for rec in records:
                ok = fetch_and_extract(rec, files_dir)
                if ok:
                    extracted_count += 1
            logger.info("  Text extracted: %d/%d", extracted_count, len(records))

        if records:
            _append_csv(csv_path, records, with_text=extract_text)
            _append_jsonl(jsonl_path, records, with_text=extract_text)

        total_docs += len(records)
        total_extracted += extracted_count
        checkpoint.mark_done(inst_id, len(records), extracted_count)

        logger.info("  Saved %d records | Running total: %d",
                    len(records), total_docs)

    # ── Summary ────────────────────────────────────────────────
    logger.info("=" * 60)
    logger.info("DONE")
    logger.info("Institutions processed : %d", len(institutions))
    logger.info("Documents collected   : %d", total_docs)
    if extract_text:
        logger.info("Text extracted        : %d", total_extracted)
    logger.info("Metadata CSV   → %s", csv_path)
    logger.info("Metadata JSONL → %s", jsonl_path)
    logger.info("=" * 60)


# ─────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────
def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description=(
            "Scrape document metadata and full text from "
            "https://www.transparencia.gob.sv/ since 2015."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    p.add_argument(
        "--extract-text", action="store_true",
        help=(
            "Download each file and extract its text content into the "
            "metadata CSV/JSONL. Supports PDF, DOCX, XLSX, XLS, HTML, TXT. "
            "Requires: pip install pdfminer.six python-docx openpyxl xlrd"
        ),
    )
    p.add_argument(
        "--download", action="store_true",
        help="Also save raw files to disk under output/transparencia/files/",
    )
    p.add_argument(
        "--id-tipo", type=int, default=None, metavar="N",
        help=(
            "Restrict to one institution category: "
            "1=Ministerios 2=Autónomas 3=Presidencia 4=Otras 5=Municipalidades "
            "6=Hospitales 7=Gobernaciones 8=ONG 10=Nuevas Munic. 11=Legislativo"
        ),
    )
    p.add_argument(
        "--institution", type=int, default=None, metavar="ID",
        help="Scrape only one institution by its numeric id_institucion",
    )
    p.add_argument(
        "--resume", action="store_true",
        help="Resume from checkpoint (skip already-completed institutions)",
    )
    p.add_argument(
        "--output-dir", type=Path, default=Path("output/transparencia"),
        help="Root output directory (default: output/transparencia/)",
    )
    p.add_argument(
        "--delay", type=float, default=1.5,
        help="Seconds between requests (default: 1.5)",
    )
    return p.parse_args()


if __name__ == "__main__":
    args = _parse_args()

    _request_delay = args.delay

    _setup_logging(args.output_dir / "logs")
    _session = _build_session()

    run(
        only_id_tipo=args.id_tipo,
        only_institution=args.institution,
        extract_text=args.extract_text,
        do_download=args.download,
        resume=args.resume,
        output_dir=args.output_dir,
    )
